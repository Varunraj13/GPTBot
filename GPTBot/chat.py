import openai
import streamlit as st
import streamlit.components.v1 as components
import pyperclip
import smtplib
import nltk
nltk.download('stopwords')
import pandas as pd
import matplotlib.pyplot as plt
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from string import punctuation
from heapq import nlargest
from email.message import EmailMessage
from io import BytesIO
import base64
import pandas as pd
import matplotlib.pyplot as plt
from fpdf import FPDF
from docx import Document
import codecs


openai.api_key = "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"

messages = [{"role": "system", "content": "Stratacent's ChatBot"}]

def CustomChatGPT(Enter_Your_Query):
    messages.append({"role": "user", "content": Enter_Your_Query})
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    ChatGPT_reply = response["choices"][0]["message"]["content"]
    messages.append({"role": "assistant", "content": ChatGPT_reply})
    return ChatGPT_reply

def send_email(receiver_email, subject, body):
    # Create a new EmailMessage object
    message = EmailMessage()

    # Set the sender and receiver email addresses
    message["From"] = "varun.sinha@stratacent.com"
    message["To"] = receiver_email

    # Set the subject and body of the email
    message["Subject"] = subject
    message.set_content(body)

    # Send the email using an SMTP server
    with smtplib.SMTP("smtp.example.com", 587) as server:
        server.login("varun.sinha@stratacent.com", "Varun@123")
        server.send_message(message)

def analyze_dataset(dataset, operation):
    # Perform analysis based on the given operation
    if operation.startswith("summarize"):
        operation_parts = operation.split(" ")
        if len(operation_parts) == 3 and operation_parts[1].isdigit() and operation_parts[2] == "lines":
            num_lines = int(operation_parts[1])
            summary = summarize_dataset(dataset, num_lines)
            analysis_result = summary if summary else "Unable to generate a summary."
        else:
            analysis_result = "Invalid operation format. Please use 'summarize X lines'."

    # elif operation.startswith("report"):
    #     operation_parts = operation.split(" ")
    #     if len(operation_parts) == 2:
    #         column_name = operation_parts[1]
    #         report = generate_pdf_report(dataset, column_name)
    #         analysis_result = report if report else f"Unable to generate a report for column '{column_name}'."
            # else:
            # analysis_result = "Invalid operation format. Please use 'report column_name'."

    elif operation.startswith("trend"):
        operation_parts = operation.split(" ")
        if len(operation_parts) == 2:
            column_name = operation_parts[1]
            trend = generate_trend(dataset, column_name)
            analysis_result = trend if trend else f"Unable to generate a trend analysis for column '{column_name}'."
        else:
            analysis_result = "Invalid operation format. Please use 'trend column_name'."

    else:
        analysis_result = "Invalid operation"

    return analysis_result


def summarize_dataset(dataset, num_lines):
    # Tokenize the dataset into sentences
    sentences = sent_tokenize(dataset)

    # Remove stopwords and punctuation
    stop_words = set(stopwords.words('english') + list(punctuation))

    # Calculate the word frequency of each word in the dataset
    word_frequencies = {}
    for word in word_tokenize(dataset.lower()):
        if word not in stop_words:
            if word not in word_frequencies:
                word_frequencies[word] = 1
            else:
                word_frequencies[word] += 1

    # Calculate the weighted frequencies of each sentence
    sentence_scores = {}
    for sentence in sentences:
        for word in word_tokenize(sentence.lower()):
            if word in word_frequencies:
                if sentence not in sentence_scores:
                    sentence_scores[sentence] = word_frequencies[word]
                else:
                    sentence_scores[sentence] += word_frequencies[word]

    # Get the top sentences based on their scores
    num_sentences = min(num_lines, len(sentences))
    summarized_sentences = nlargest(num_sentences, sentence_scores, key=sentence_scores.get)

    # Combine the summarized sentences into a summary
    summary = " ".join(summarized_sentences)

    return summary

def read_csv_file(file_path):
    try:
        df = pd.read_csv(file_path)
        return df
    except Exception as e:
        return None

def filter_data(df, condition):
    try:
        filtered_df = df[eval(condition)]
        return filtered_df
    except Exception as e:
        return None

def sort_data(df, sort_by):
    try:
        sorted_df = df.sort_values(by=sort_by)
        return sorted_df
    except Exception as e:
        return None

def aggregate_data(df, group_by, aggregation):
    try:
        aggregated_df = df.groupby(group_by).agg(aggregation)
        return aggregated_df
    except Exception as e:
        return None

def modify_data(df, column, new_value):
    try:
        df[column] = new_value
        return df
    except Exception as e:
        return None

def merge_data(df1, df2, common_column):
    try:
        merged_df = pd.merge(df1, df2, on=common_column)
        return merged_df
    except Exception as e:
        return None

def export_csv(df, file_path):
    try:
        df.to_csv(file_path, index=False)
        return "CSV file exported successfully."
    except Exception as e:
        return "Error occurred while exporting the CSV file."

def generate_trend(dataset, column_name):
    try:
        df = pd.read_csv(dataset)
        if column_name in df.columns and df[column_name].dtype in [int, float]:
            df.plot(x="date", y=column_name)
            plt.xlabel("Date")
            plt.ylabel(column_name)
            plt.title(f"Trend Analysis for '{column_name}'")
            plt.show()
            trend = "Trend plot generated."
        else:
            trend = f"Invalid column '{column_name}' or it is not numerical."

    except Exception as e:
        trend = "Error occurred while generating the trend analysis."

    return trend

def generate_download_link(content, file_format):
    if file_format == "PDF":
        buffer = BytesIO()
        content.save(buffer, format="PDF")
        buffer.seek(0)
        b64 = base64.b64encode(buffer.read()).decode()
        href = f'<a href="data:application/pdf;base64,{b64}" download="response.pdf">Download PDF</a>'
    elif file_format == "Word":
        buffer = BytesIO()
        content.save(buffer)
        buffer.seek(0)
        b64 = base64.b64encode(buffer.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="response.docx">Download Word</a>'
    else:
        href = ""
    return href




def clear_messages():
    session_state.response = ""
    session_state.analysis_result = ""

session_state = st.session_state
if "response" not in session_state:
    session_state.response = ""

title = "StrataBot 🤖"
description = "Chat with StrataBot"

prompts = [
    "Write 3 Introductory Emails for an Organization",
    "3 Templates For a Financial Report",
    "Pointers for a Product Presentation",
    "Social media sample for Hiring",
    "3 Leave Applications to a Manager",
    "3 Referral Emails for an Organization",
    "3 Sample Sales pointers for a PPT",
    "3 different Briefings about a company",
    "Pointers for a Sample Analytical Report",
    "Company Profile of Stratacent",
    "Email for Work from Home",
    "Event Announcement Email"
]

st.set_page_config(page_title=title, page_icon="🤖")

# Custom CSS styles
custom_styles = """
<style>
body {
  margin: 0;
  padding: 0;
  background-image: url('https://media.glassdoor.com/banner/bh/1186570/stratacent-banner-1646985658882.jpg');
  background-repeat: round;
  background-size: cover;
}

#main-container {
  padding: 20px;
  background-color: black;
  min-height: 100vh;
}

#chatbot-container {
  background-color: white;
  padding: 20px;
  border-radius: 30px;
  margin-top: 20px;
  max-width: 100%;
}

.light-blue-button {
  background-color: #C7F6C7;
  border: none;
  color: black;
  padding: 10px 20px;
  text-align: center;
  text-decoration: none;
  display: inline-block;
  font-size: 16px;
  margin: 4px 2px;
  cursor: pointer;
}

.button-row {
  display: flex;
  justify-content: center;
  margin-top: 10px;
}

.visit-stratacent-button {
  background-color: #C7F6C7;
  border: none;
  color: black;
  padding: 10px 20px;
  text-align: center;
  text-decoration: none;
  display: inline-block;
  font-size: 16px;
  margin-top: 20px;
  cursor: pointer;
}
</style>
"""

# Apply custom CSS styles
components.html(custom_styles)

# Render the page
st.markdown(
    f'<div id="main-container"><h1 style="color: orange;">{title}</h1><p style="color: blue;">{description}</p></div>',
    unsafe_allow_html=True
)

user_input = st.text_input("Enter your query")

if st.button("Ask", key="user_input_button", help="light-blue-button"):
    if user_input:
        ChatGPT_reply = CustomChatGPT(user_input)
        st.text_area("Response", ChatGPT_reply, height=200, key="response_area", help="chatbot-container")

        st.subheader("Was the Answer Useful?")
        yes_button = st.button("Yes", key="yes_button", help="light-blue-button")
        no_button = st.button("No", key="no_button", help="light-blue-button")
        if yes_button or no_button:
            st.write("Thank you for your response.")
            clear_messages()

        file_format = st.selectbox("Select the file format for download", ["PDF", "Word"])
        download_button = st.button("Download", key="download_button", help="light-blue-button")

        if download_button:
            if file_format == "PDF":
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, ChatGPT_reply)
                pdf_output = BytesIO()
                pdf.output(pdf_output)
                pdf_bytes = pdf_output.getvalue()
                b64_pdf = base64.b64encode(pdf_bytes).decode("latin-1")
                href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="response.pdf">Download PDF</a>'
                st.markdown(href, unsafe_allow_html=True)
            elif file_format == "Word":
                doc = Document()
                doc.add_paragraph(ChatGPT_reply)
                doc_output = BytesIO()
                doc.save(doc_output)
                doc_bytes = doc_output.getvalue()
                b64_doc = base64.b64encode(doc_bytes).decode("latin-1")
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_doc}" download="response.docx">Download Word</a>'
                st.markdown(href, unsafe_allow_html=True)
        # st.subheader("Send the Response via Email")
        # receiver_email = st.text_input("Receiver's Email")
        # if st.button("Send Email", key="send_email_button", help="light-blue-button"):
        #     if receiver_email:
        #         subject = "Response from StrataBot"
        #         body = ChatGPT_reply
        #         send_email(receiver_email, subject, body)
        #         st.write(f"Email sent to {receiver_email}).")

st.subheader("Choose a prompt:")

col1, col2, col3 = st.columns(3)

for i, prompt in enumerate(prompts):
    if i % 3 == 0:
        col = col1
    elif i % 3 == 1:
        col = col2
    else:
        col = col3

    if col.button(prompt, key=f"prompt_button_{i}", help="light-blue-button"):
        ChatGPT_reply = CustomChatGPT(prompt)
        st.text_area("Response", ChatGPT_reply, height=750, key="response_area", help="chatbot-container")
        st.subheader("Was the Answer Useful?")
        yes_button = st.button("Yes", key="yes_button", help="light-blue-button")
        no_button = st.button("No", key="no_button", help="light-blue-button")

        if yes_button or no_button:
            st.write("Thank you for your response.")
            if yes_button:
                # Perform action if the answer was useful
                pass
            elif no_button:
                # Perform action if the answer was not useful
                pass

        # if st.button("Copy to Clipboard", key="copy_button", help="light-blue-button"):
        #     pyperclip.copy(ChatGPT_reply)
        #     st.write("Response copied to clipboard.")

        file_format = st.selectbox("Select the file format for download", ["PDF", "Word"])
        download_button = st.button("Download", key="download_button", help="light-blue-button")

        if download_button:
            ChatGPT_reply = st.text_area("Response", height=750, key="response_area", help="chatbot-container")
            if file_format == "PDF":
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, ChatGPT_reply)
                pdf_output = BytesIO()
                pdf.output(pdf_output)
                pdf_bytes = pdf_output.getvalue()
                b64_pdf = base64.b64encode(pdf_bytes).decode()
                href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="response.pdf">Download PDF</a>'
                st.markdown(href, unsafe_allow_html=True)
            elif file_format == "Word":
                doc = Document()
                doc.add_paragraph(ChatGPT_reply)
                doc_output = BytesIO()
                doc.save(doc_output)
                doc_bytes = doc_output.getvalue()
                b64_doc = base64.b64encode(doc_bytes).decode()
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_doc}" download="response.docx">Download Word</a>'
                st.markdown(href, unsafe_allow_html=True)

uploaded_file = st.file_uploader("Upload a file", type=["csv", "txt"])

if uploaded_file is not None:
    file_type = uploaded_file.type

    if file_type == "text/plain":
        # Perform text summarization
        text = uploaded_file.read().decode("utf-8")
        num_lines = st.number_input("Enter the number of lines for the summary", min_value=1, max_value=100, value=3)
        summary = summarize_dataset(text, num_lines)

        # Display the generated summary
        st.subheader("Summary")
        st.write(summary)

    elif file_type == "text/csv":
        # Read the uploaded CSV file
        df = pd.read_csv(uploaded_file)

        # Display the uploaded data
        st.subheader("Uploaded Data")
        st.write(df)

        # Perform operations on the uploaded data
        operation = st.text_input("Enter the operation (e.g., trend sales):")

        if st.button("Analyze"):
            operation_parts = operation.split(" ")
            if len(operation_parts) == 2:
                operation_type = operation_parts[0]
                column_name = operation_parts[1]

                if operation_type == "trend":
                    trend_analysis = generate_trend(df, column_name)
                    st.subheader("Trend Analysis")
                    st.write(trend_analysis)
                else:
                    st.write("Invalid operation. Please enter a valid operation.")

            else:
                st.write("Invalid operation. Please enter a valid operation.")

    file_format = st.selectbox("Select the file format for download", ["PDF", "Word"])
    download_button = st.button("Download", key="download_button", help="light-blue-button")
    if download_button:
        if file_format == "PDF":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            encoded_summary = codecs.encode(summary, 'utf-8').decode('latin-1')
            pdf.multi_cell(0, 10, encoded_summary)

        # Generate the download link
        pdf_output = pdf.output(dest="S").encode("latin-1")
        base64_pdf = base64.b64encode(pdf_output).decode("utf-8")
        download_link = f'<a href="data:application/pdf;base64,{base64_pdf}" download="summary.pdf">Download PDF</a>'

    elif file_format == "Word":
        doc = Document()
        encoded_summary = codecs.encode(summary, 'utf-8').decode('latin-1')
        doc.add_paragraph(encoded_summary)
        doc_output = BytesIO()
        doc.save(doc_output)

        # Generate the download link
        base64_doc = base64.b64encode(doc_output.getvalue()).decode("utf-8")
        download_link = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64_doc}" download="summary.docx">Download Word</a>'

    else:
        download_link = ""
    st.markdown(download_link, unsafe_allow_html=True)
        # st.subheader("Send the Response via Email")
        # receiver_email = st.text_input("Receiver's Email")
        # if st.button("Send Email", key="send_email_button", help="light-blue-button"):
        #     if receiver_email:
        #         subject = "Response from StrataBot"
        #         body = report
        #         send_email(receiver_email, subject, body)
        #         st.write(f"Email sent to {receiver_email}.")

st.subheader("About Stratacent")
st.write("Stratacent is a leading technology consulting and solutions company with a focus on the financial services industry. "
         "We provide innovative solutions and services that help our clients achieve their business goals and stay ahead in a rapidly evolving market.")
st.write("Visit [Stratacent](https://www.stratacent.com) for more information.")